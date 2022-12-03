# https://pypi.org/project/python-docx/
# https://python-docx.readthedocs.io/en/latest/
# pip install python-docx
from docx           import Document

from docx.shared    import Pt

from docx.shared    import RGBColor

from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

# ******************************************************
#                   PAGE 1 FUNCTIONS
# ******************************************************
def addTitle_P1(document):

    # print(dir(document))
    # print(help(document.add_heading))

    # title = document.add_heading('Koala', level= 0)
    title = document.add_heading('Koala', level= 1)
    # title = document.add_paragraph('Koala')


    # print(type(title))
    # print(dir(title))
    # print(dir(title.style))
    # print(dir(title.style.font))
    title.style.font.name = 'Calibri (Body)'


    # print(dir(title.style.font))
    # title.style.font.size = '26'
    # title.style.font.size = 26
    # print(type(title.style.font.size))
    # print(help(Pt))
    title.style.font.size = Pt(26)


    # print(dir(title.style.font))
    # print(type(title.style.font.bold))
    # print(title.style.font.bold)
    title.style.font.bold = False
    # title.style.font.italic = True
    # title.style.font.underline = True
    # print(type(title.style.font.underline))
    # print(type(title.style.font))


    # print(dir(title.style.font))
    # title.style.font.color = 'Red'
    # print(type(title.style.font.color))
    # print(dir(title.style.font.color))
    # print(type(title.style.font.color.rgb))
    # print(help(RGBColor))
    title.style.font.color.rgb = RGBColor(46, 116, 181)


    # print(dir(title))
    # title.alignment = 'center'
    # print(type(title.alignment))
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    ...
# END def addTitle_P1

def addParagraph_P1(document):

    ...
# END def addParagraph_P1

def addKoalaImage_P1(document):

    ...
# END def addKoalaImage_P1


# ******************************************************
#                   PAGE 2 FUNCTIONS
# ******************************************************
def addTitle_P2(document):

    ...
# END addTitle_P2

def addTable_P2(document):

    ...
# END addTable_P2


# ******************************************************
#                   DOCUMENT CREATION
# ******************************************************
document = Document()

# PAGE 1
addTitle_P1(document)
addParagraph_P1(document)
addKoalaImage_P1(document)

# PAGE 2
addTitle_P2(document)
addTable_P2(document)

document.save('koala.docx')